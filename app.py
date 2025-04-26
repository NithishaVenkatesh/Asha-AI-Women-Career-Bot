from flask import Flask, render_template, request, jsonify
from asha import AshaBot
import json
import re
import os
import mimetypes
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import tempfile
from langdetect import detect

app = Flask(__name__)
bot = AshaBot()

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create uploads folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_type(filename):
    """Determine file type from extension"""
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    if ext == 'pdf':
        return 'application/pdf'
    elif ext in ['doc', 'docx']:
        return 'application/msword'
    return ''

def try_ocr(file_path):
    """Attempt OCR if available"""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        
        text = ""
        images = convert_from_path(file_path)
        for image in images:
            text += pytesseract.image_to_string(image) + "\n"
        return text.strip()
    except Exception as e:
        print(f"OCR failed: {str(e)}")
        return ""

def extract_text_from_pdf(file_path):
    """Extract text from PDF using PyPDF2 and OCR if needed"""
    text = ""
    
    # First try normal PDF text extraction
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        print(f"Error in PDF text extraction: {str(e)}")
    
    # If text extraction yields little text, try OCR
    if len(text.strip()) < 100:  # Arbitrary threshold for "too little text"
        ocr_text = try_ocr(file_path)
        if ocr_text:
            text = ocr_text
    
    return text.strip()

def extract_text_from_docx(file_path):
    """Extract text from DOCX with improved formatting"""
    text = []
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
    except Exception as e:
        print(f"Error in DOCX extraction: {str(e)}")
        raise
    
    return "\n".join(text)

def process_resume_text(text):
    """Process and structure the resume text before sending to bot"""
    if not text.strip():
        raise ValueError("No text could be extracted from the document")
        
    # Clean up the text
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces with single space
    text = re.sub(r'\n\s*\n', '\n', text)  # Remove multiple empty lines
    
    # Structure the request to the bot
    prompt = (
        "I have received a resume to analyze. Please extract and organize the following information:\n"
        "1. Contact Information\n"
        "2. Education History\n"
        "3. Work Experience\n"
        "4. Skills (both technical and soft skills)\n"
        "5. Projects (if any)\n"
        "6. Certifications (if any)\n\n"
        "Here is the resume content:\n\n"
        f"{text}\n\n"
        "Please analyze this information and provide a structured summary. "
        "Also highlight any notable achievements or unique qualifications."
    )
    return prompt

@app.route('/')
def welcome():
    return render_template('welcome.html')

@app.route('/chat')
def chat_page():
    return render_template('chat.html')

@app.route('/upload-resume', methods=['POST'])
def upload_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['resume']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        try:
            # Create a temporary file to store the upload
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                file.save(temp_file.name)
                
                try:
                    # Get file type from extension
                    file_type = get_file_type(file.filename)
                    
                    # Extract text based on file type
                    if 'pdf' in file_type:
                        text = extract_text_from_pdf(temp_file.name)
                    elif 'msword' in file_type or file.filename.endswith(('.docx', '.doc')):
                        text = extract_text_from_docx(temp_file.name)
                    else:
                        return jsonify({'error': 'Invalid file type'}), 400
                    
                    # Process the text and create a structured prompt
                    prompt = process_resume_text(text)
                    
                    # Send to bot for analysis
                    try:
                        response = bot.chat.send_message(prompt)
                        return jsonify({
                            'response': response.text,
                            'success': True
                        })
                    except Exception as e:
                        return jsonify({
                            'error': f"Error processing resume: {str(e)}",
                            'success': False
                        }), 500
                        
                finally:
                    # Clean up the temporary file
                    try:
                        os.unlink(temp_file.name)
                    except:
                        pass
                    
        except Exception as e:
            return jsonify({
                'error': f"Error processing file: {str(e)}",
                'success': False
            }), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/chat', methods=['POST'])
def chat():
    user_input = request.json.get('message', '')
    if user_input.lower() in ["exit", "quit", "bye"]:
        return jsonify({
            'response': "It was wonderful chatting with you! 💖 Wishing you all the very best on your career journey! Stay shining! 🌟",
            'exit': True
        })
    
    try:
        response = bot.chat.send_message(user_input)
        response_text = response.text.lower()
        
        # Check if this is a job search confirmation
        confirmation_phrases = [
            "i will search", "let me look for", "searching for",
            "look for some opportunities", "find opportunities",
            "search right away", "search now"
        ]
        
        if any(phrase in response_text for phrase in confirmation_phrases) and "jobs" in response_text:
            # Extract skills and location from the response
            skills_match = re.search(r'skills?[:\s\'"]+([\w\s,-]+)[\'"]?', response_text)
            location_match = re.search(r'(?:location|in|at|for)[:\s\'"]+([\w\s,-]+)[\'"]?', response_text)
            
            if skills_match and location_match:
                skills = skills_match.group(1).strip()
                location = location_match.group(1).strip()
                
                if location != 'location':  # Avoid matching the word "location" itself
                    print(f"Initiating job search with skills='{skills}', location='{location}'")
                    try:
                        jobs = bot.search_jobs_external(skills, location)
                        
                        if jobs:
                            # Format job results in a user-friendly way
                            job_response = ["Here are some opportunities I found for you:\n"]
                            for i, job in enumerate(jobs, 1):
                                job_response.append(f"{i}. **{job.get('title', 'N/A')}**")
                                job_response.append(f"   🏢 Company: {job.get('company', 'N/A')}")
                                job_response.append(f"   📍 Location: {job.get('location', 'N/A')}")
                                job_response.append(f"   🔗 [View Job]({job.get('link', 'N/A')})")
                                job_response.append(f"   ℹ️ Source: {job.get('source', 'N/A')}")
                                job_response.append(f"   💯 Match: {job.get('match_percentage', 0):.0f}% match with your skills\n")
                            
                            return jsonify({'response': '\n'.join(job_response)})
                    except Exception as scraper_error:
                        print(f"Error during job search: {str(scraper_error)}")
                        return jsonify({'response': "I encountered an error while searching for jobs. Would you like to try again?"})
        
        return jsonify({
            'response': response.text,
            'exit': False
        })
    except Exception as e:
        return jsonify({
            'response': f"Oh dear, I seem to have encountered a little technical hiccup. Could you perhaps try phrasing that differently? My apologies! 😥 ({str(e)})",
            'exit': False
        })

@app.route('/detect-language', methods=['POST'])
def detect_language():
    try:
        text = request.json.get('text', '')
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        # Detect language
        language = detect(text)
        
        # Map language code to BCP-47 format for Web Speech API
        language_map = {
            'en': 'en-US',
            'es': 'es-ES',
            'fr': 'fr-FR',
            'de': 'de-DE',
            'it': 'it-IT',
            'pt': 'pt-BR',
            'ru': 'ru-RU',
            'ja': 'ja-JP',
            'ko': 'ko-KR',
            'zh': 'zh-CN',
            'hi': 'hi-IN',
            'ar': 'ar-SA'
        }
        
        # Return the mapped language code or default to en-US
        return jsonify({
            'language': language_map.get(language, 'en-US')
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))  # get Railway-assigned port
    app.run(host="0.0.0.0", port=port)
